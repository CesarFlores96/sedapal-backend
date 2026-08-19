import io
import json
import zipfile
import datetime
from typing import Any
import pandas as pd
from docx import Document
from docx.shared import Pt
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, Body
from fastapi.responses import JSONResponse, Response
from pathlib import Path
from psycopg.rows import dict_row
import traceback

from app.database import get_pool

router = APIRouter(prefix="/billing/credit-notes", tags=["Credit Notes"])


def clean_val(val: Any) -> str:
    if pd.isna(val) or val == "NaN":
        return ""
    if isinstance(val, (pd.Timestamp, datetime.datetime)):
        meses_abrev = ["Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
        return f"{meses_abrev[val.month - 1]}-{str(val.year)[-2:]}"
    return str(val).strip()


@router.post("/preview")
async def preview_credit_notes(
    file: UploadFile = File(...),
    pool=Depends(get_pool),
) -> JSONResponse:
    try:
        if not file.filename.endswith(".xlsx"):
            raise HTTPException(status_code=400, detail="El archivo debe ser un .xlsx")

        contents = await file.read()
        
        # Asumimos que los encabezados están en la fila 1 (0 indexado) o saltamos vacíos
        df = pd.read_excel(io.BytesIO(contents))
        
        # A veces el excel de sedapal tiene el header en la segunda fila, busquemos donde está "CU"
        header_idx = None
        for i, row in df.iterrows():
            row_str = " ".join([str(x) for x in row.values])
            if "CU" in row_str and "CLIENTE" in row_str:
                header_idx = i
                break
                
        if header_idx is not None:
            df = pd.read_excel(io.BytesIO(contents), header=header_idx + 1)
            
        # Si las columnas no se detectan, intentamos un fallback
        if "CU" not in df.columns:
            df = pd.read_excel(io.BytesIO(contents))
            # Buscar manualmente

        results = []
        
        async with pool.connection() as conn:
            async with conn.cursor(row_factory=dict_row) as cursor:
                for idx, row in df.iterrows():
                    cu = clean_val(row.get("CU"))
                    if not cu or cu == "CU":
                        continue
                        
                    # Buscar info en DB
                    query_supply = """
                        SELECT customer_name, service_address, district
                        FROM public.customer_supplies
                        WHERE supply_code = %s
                        LIMIT 1
                    """
                    await cursor.execute(query_supply, (cu,))
                    supply = await cursor.fetchone()
                    
                    customer_name = supply["customer_name"] if supply else clean_val(row.get("CLIENTE"))
                    address = supply["service_address"] if supply else clean_val(row.get("DIRECCIÓN ENVIO"))
                    district = supply["district"] if supply else clean_val(row.get("DISTRITO"))
                    
                    # Buscar último consumo y facturación
                    query_billing = """
                        SELECT billed_volume_m3, total_soles, period_year, period_month
                        FROM public.customer_debts
                        WHERE supply_code = %s
                        ORDER BY period_year DESC, period_month DESC
                        LIMIT 1
                    """
                    await cursor.execute(query_billing, (cu,))
                    billing = await cursor.fetchone()
                    last_consumption = float(billing["billed_volume_m3"]) if billing and billing["billed_volume_m3"] else 0.0
                    last_billing = float(billing["total_soles"]) if billing and billing["total_soles"] else 0.0
                
                    res_row = {
                        "id": str(idx),
                        "CU": cu,
                        "CARTA": clean_val(row.get("CARTA")),
                        "CLIENTE": customer_name,
                        "DIRECCION_ENVIO": address,
                        "DISTRITO": district,
                        "CONCEPTO": clean_val(row.get("CONCEPTO")),
                        "MES_ANO": clean_val(row.get("MES/AÑO")),
                        "MEDIDOR": clean_val(row.get("MEDIDOR")),
                        "VOL_EMITIDO": clean_val(row.get("VOL EMITIDO")),
                        "IMPORT_EMITIDO": clean_val(row.get("IMPORT EMITIDO")),
                        "VOL_CORREGIDO": clean_val(row.get("VOL CORRREGIDO")),
                        "IMPORT_CORREGIDO": clean_val(row.get("IMPORT CORREGIDO")),
                        "IMPORT_COMPROBANTE": clean_val(row.get("IMPORT COMPROBANTE")),
                        "NUMERO_COMPROBANTE": clean_val(row.get("NUMERO COMPROBANTE")),
                        "TOTAL_DEL_RECIBO": clean_val(row.get("TOTAL DEL RECIBO")),
                        "consumo_ultimo_mes": last_consumption,
                        "facturacion_ultimo_mes": last_billing
                    }
                    results.append(res_row)
                
        return JSONResponse(content={"success": True, "data": results})

    except Exception as e:
        error_details = traceback.format_exc()
        # Imprimimos en consola para backend logs
        print("ERROR EN PREVIEW:", error_details)
        return JSONResponse(status_code=500, content={"success": False, "error": str(e), "traceback": error_details})


def replace_text_in_paragraph(paragraph, mapping):
    for key, val in mapping.items():
        key_str = str(key)
        val_str = str(val)
        if key_str in paragraph.text:
            inline = paragraph.runs
            # Intentar reemplazar dentro de un mismo run
            for i in range(len(inline)):
                if key_str in inline[i].text:
                    inline[i].text = inline[i].text.replace(key_str, val_str)
            
            # Si la palabra clave todavía está en el párrafo, significa que está dividida en múltiples runs.
            # En este caso preservamos el formato del primer run y borramos el resto.
            if key_str in paragraph.text:
                full_text = paragraph.text.replace(key_str, val_str)
                if len(inline) > 0:
                    inline[0].text = full_text
                    for i in range(1, len(inline)):
                        inline[i].text = ""

def get_current_date_es():
    meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    hoy = datetime.datetime.now()
    return f"Lima, {hoy.day} de {meses[hoy.month - 1]} de {hoy.year}"

def get_month_from_mesano(mes_ano: str) -> str:
    meses_map = {
        "ene": "enero", "feb": "febrero", "mar": "marzo", "abr": "abril", "may": "mayo", "jun": "junio",
        "jul": "julio", "ago": "agosto", "sep": "septiembre", "oct": "octubre", "nov": "noviembre", "dic": "diciembre"
    }
    if not mes_ano:
        return ""
    prefix = str(mes_ano).split("-")[0].lower()
    return meses_map.get(prefix, str(mes_ano))

@router.post("/generate")
async def generate_credit_notes(
    data: str = Form(...),
):
    """
    data: JSON string de los registros seleccionados
    """
    try:
        records = json.loads(data)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON data")

    template_path = Path(__file__).resolve().parent.parent / "templates" / "Carta.docx"
    if not template_path.exists():
        raise HTTPException(status_code=500, detail="Plantilla no encontrada en el servidor.")
        
    with open(template_path, "rb") as f:
        template_bytes = f.read()
    
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for record in records:
            # Crear un nuevo doc desde la plantilla
            doc = Document(io.BytesIO(template_bytes))
            
            # Mapeo de variables
            mapping = {
                "«CARTA»": record.get("CARTA", ""),
                "«CLIENTE»": record.get("CLIENTE", ""),
                "«DIRECCIÓN_ENVIO»": record.get("DIRECCION_ENVIO", ""),
                "«DISTRITO»": record.get("DISTRITO", ""),
                "«CU»": record.get("CU", ""),
                "<<Asunto>>": record.get("CONCEPTO", ""),
                "«MESAÑO»": record.get("MES_ANO", ""),
                "<<Mes>>": get_month_from_mesano(record.get("MES_ANO", "")),
                "«TOTAL_DEL_RECIBO»": str(record.get("TOTAL_DEL_RECIBO", "")),
                "«NUMERO_COMPROBANTE_»": record.get("NUMERO_COMPROBANTE", ""),
                "«NUMERO_COMPROBANTE»": record.get("NUMERO_COMPROBANTE", ""),
                "«MEDIDOR»": record.get("MEDIDOR", ""),
                "«VOL_EMITIDO»": str(record.get("consumo_ultimo_mes", "")),
                "«IMPORT_EMITIDO»": str(record.get("facturacion_ultimo_mes", "")),
                "«VOL_CORRREGIDO»": record.get("VOL_CORREGIDO", ""),
                "«IMPORT_CORREGIDO»": record.get("IMPORT_CORREGIDO", ""),
                "«IMPORT_COMPROBANTE»": record.get("IMPORT_COMPROBANTE", ""),
                "Lima, 17 de agosto de 2026": get_current_date_es(),
                "«FECHA»": get_current_date_es(),
            }
            
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, mapping)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, mapping)
                            
            # Forzar formato exacto solicitado
            for p in doc.paragraphs:
                text_lower = p.text.lower()
                size = Pt(10)
                if "año de la esperanza" in text_lower or "decenio de la igualdad" in text_lower:
                    size = Pt(8)
                for run in p.runs:
                    if run.text.strip():
                        run.font.name = 'Tahoma'
                        run.font.size = size

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if run.text.strip():
                                    run.font.name = 'Tahoma'
                                    run.font.size = Pt(7.5)
                        
            # Guardar el docx en memoria
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            filename = f"Carta_NC_{record.get('CU', 'desc')}.docx"
            zip_file.writestr(filename, doc_io.read())
            
        # Además agregamos el Excel combinado
        df_out = pd.DataFrame(records)
        excel_io = io.BytesIO()
        with pd.ExcelWriter(excel_io, engine="openpyxl") as writer:
            df_out.to_excel(writer, index=False)
        excel_io.seek(0)
        
        zip_file.writestr("Resumen_Cartas_Generadas.xlsx", excel_io.read())
        
    zip_buffer.seek(0)
    
    return Response(
        content=zip_buffer.read(),
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=cartas_generadas.zip"}
    )
